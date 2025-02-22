from django.shortcuts import render, redirect
from django.http import HttpResponse, FileResponse
from .models import PastCase
from .forms import SolicitorForm, WhoAreYouReportingForm, JCIOForm,  BarristerForm, NewspaperForm, BankForm
from django.db.models import Q
import os
from django.http import HttpResponse
from docx import Document
from django.conf import settings
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .serializers import SubmissionSerializer
from django.core.mail import send_mail
from rest_framework.views import APIView
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from rest_framework import status
from .serializers import DocumentSerializer, QuerySerializer
from .models import Document
from .utils import DocumentProcessor, QueryProcessor
import openai
from django.conf import settings

def who_are_you_reporting(request):
    if request.method == 'POST':
        form = WhoAreYouReportingForm(request.POST)
        if form.is_valid():
            report_type = form.cleaned_data['report_type']
            
            # Redirect to the appropriate form based on user selection
            if report_type == 'solicitor':
                return redirect('solicitor_form')  # URL for solicitor (SRA) form
            elif report_type == 'barrister':
                return redirect('barrister_form')  # URL for barrister (BSB) form
            elif report_type == 'judge':
                return redirect('judge_form')  # URL for judge form
            elif report_type == 'newspaper':
                return redirect('newspaper_form')
            elif report_type == 'bank':
                return redirect('bank_form')
    else:
        form = WhoAreYouReportingForm()
    
    return render(request, 'complaints/who_are_you_reporting.html', {'form': form})


def find_similar_cases(complaint_category, complaint_text):
    keywords = complaint_text.split()
    query = PastCase.objects.filter(category=complaint_category)
    
    q = Q()
    for keyword in keywords:
        q |= Q(keywords__icontains=keyword)
    
    query = query.filter(q)
    return query

def check_for_breaches(complaint_text):
    breaches = []
    if "CPR" in complaint_text:
        breaches.append("CPR Breach")
    if "human rights" in complaint_text:
        breaches.append("Human Rights Breach")
    if "rule of law" in complaint_text:
        breaches.append("Rule of Law Breach")
    return breaches

def populate_complaint_form(complaint_category, form_data, similar_cases):
    if complaint_category == "solicitor":
        template_path = os.path.join(settings.BASE_DIR, 'templates/forms/sra.docx')
    elif complaint_category == "barrister":
        template_path = os.path.join(settings.BASE_DIR, 'templates/forms/bsb.docx')
    elif complaint_category == "judge":
        template_path = os.path.join(settings.BASE_DIR, 'templates/forms/jcio.docx')
    elif complaint_category == "newspaper":
        template_path = os.path.join(settings.BASE_DIR, 'templates/forms/newspaper.docx')
    elif complaint_category == "bank":
        template_path = os.path.join(settings.BASE_DIR, 'templates/forms/bank.docx')
    else:
        return None

    # Load the template
    doc = Document(template_path)

    # Function to replace placeholders in paragraphs
    def replace_text(paragraph, placeholder, value):
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, str(value))

    # Replace placeholders inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if complaint_category == "solicitor":
                        replace_solicitor_placeholders(paragraph, form_data, replace_text)
                    elif complaint_category == "barrister":
                        replace_barrister_placeholders(paragraph, form_data, replace_text)
                    elif complaint_category == "judge":
                        replace_judge_placeholders(paragraph, form_data, replace_text)
                    elif complaint_category == "newspaper":
                        replace_newspaper_placeholders(paragraph, form_data, replace_text)
                    elif complaint_category == "bank":
                        replace_bank_placeholders(paragraph, form_data, replace_text)

    # Save the populated document
    populated_form_path = os.path.join(settings.MEDIA_ROOT, f'populated_{complaint_category}_complaint_form.docx')
    doc.save(populated_form_path)

    return populated_form_path

def replace_solicitor_placeholders(paragraph, form_data, replace_text):
    replace_text(paragraph, '<<TITLE>>', form_data['title'])
    replace_text(paragraph, '<<FIRST_NAME>>', form_data['first_name'])
    replace_text(paragraph, '<<LAST_NAME>>', form_data['last_name'])
    replace_text(paragraph, '<<FIRM>>', form_data['firm'] or 'N/A')
    replace_text(paragraph, '<<ADDRESS>>', form_data['address'])
    replace_text(paragraph, '<<PHONE>>', form_data['phone_number'])
    replace_text(paragraph, '<<EMAIL>>', form_data['email'])
    replace_text(paragraph, '<<ADJUSTMENTS>>', form_data['adjustments'] or 'N/A')
    replace_text(paragraph, '<<INDIVIDUAL_ACTED_FOR>>', form_data['individual_acted_for'])
    replace_text(paragraph, '<<REPORTED_PERSON_NAME>>', form_data['reported_person_name'])
    replace_text(paragraph, '<<REPORTED_FIRM_NAME>>', form_data['reported_firm_name'])
    replace_text(paragraph, '<<REPORTED_FIRM_ADDRESS>>', form_data['reported_firm_address'])
    replace_text(paragraph, '<<REPORTED_FIRM_PHONE>>', form_data['reported_firm_phone'])
    replace_text(paragraph, '<<ACTED_FOR_YOU>>', form_data['acted_for_you'])
    replace_text(paragraph, '<<ACT_FOR_ANOTHER_PERSON>>', form_data['act_for_another_person'])
    replace_text(paragraph, '<<INDIVIDUAL_ACTING_FOR>>', form_data['individual_acting_for'] or 'N/A')
    replace_text(paragraph, '<<COMPLAINT>>', form_data['complaint'])
    replace_text(paragraph, '<<SIGNATURE>>', form_data['signature'])
    replace_text(paragraph, '<<DATE>>', form_data['date'].strftime("%d/%m/%Y"))

def replace_barrister_placeholders(paragraph, form_data, replace_text):
        # Section 1: Who you're reporting
    replace_text(paragraph, '<<BARRISTER_NAME>>', form_data['barrister_name'])
    replace_text(paragraph, '<<REPORTED_ADDRESS>>', form_data['address1'])
    replace_text(paragraph, '<<REPORTED_POSTCODE>>', form_data['postcode'])
    replace_text(paragraph, '<<REPORTED_PHONE>>', form_data['phone_number'])
    replace_text(paragraph, '<<REPORTED_EMAIL>>', form_data['email'])

    # Section 2: Reporting more than one barrister
    replace_text(paragraph, '<<MORE_THAN_ONE>>', form_data['more_than_one'])
    replace_text(paragraph, '<<BARRISTER_NAME2>>', form_data['second_barrister_name'] or 'N/A')
    replace_text(paragraph, '<<ADDRESS2>>', form_data['second_address1'] or 'N/A')
    replace_text(paragraph, '<<POSTCODE2>>', form_data['second_postcode'] or 'N/A')
    replace_text(paragraph, '<<EMAIL2>>', form_data['second_email'] or 'N/A')
    replace_text(paragraph, '<<PHONE2>>', form_data['second_phone_number'] or 'N/A')

    # Section 3: Case-related information
    replace_text(paragraph, '<<ACTING_FOR>>', form_data['acting_for'])
    replace_text(paragraph, '<<RELATED_TO_REPORT>>', form_data['related_to_previous_report'])
    replace_text(paragraph, '<<REFERENCE_NUMBER>>', form_data['reference_number'] or 'N/A')

    # Section 4: Court case details
    replace_text(paragraph, '<<RELATED_TO_COURT>>', form_data['related_to_court_case'])
    replace_text(paragraph, '<<CASE_NAME>>', form_data['case_name'] or 'N/A')
    replace_text(paragraph, '<<COURT_NAME>>', form_data['court_name'] or 'N/A')
    replace_text(paragraph, '<<COURT_REF>>', form_data['court_reference_number'] or 'N/A')

    # Section 5: Your relationship to the case
    replace_text(paragraph, '<<RELATIONSHIP_TO_CASE>>', form_data['relationship_to_case'])
    replace_text(paragraph, '<<LITIGANT_IN_PERSON>>', form_data['litigant_in_person'])
    replace_text(paragraph, '<<CASE_ONGOING>>', form_data['case_ongoing'])

    # Section 6: Event details
    replace_text(paragraph, '<<LAST_OCCURRENCE>>', form_data['last_occurrence'].strftime("%d/%m/%Y"))
    replace_text(paragraph, '<<EVENT_INFO>>', form_data['event_information'])

    # Section 7: Additional witnesses
    replace_text(paragraph, '<<WITNESS_NAME>>', form_data['witness_name'] or 'N/A')
    replace_text(paragraph, '<<WITNESS_EMAIL>>', form_data['witness_email'] or 'N/A')
    replace_text(paragraph, '<<WITNESS_PHONE>>', form_data['witness_phone_number'] or 'N/A')
    replace_text(paragraph, '<<WITNESS_CONSENT>>', form_data['witness_consent'] or 'N/A')

    # Section 8: Documentation and additional information
    replace_text(paragraph, '<<OTHER_INFO>>', form_data['other_information'] or 'N/A')

    # Section 9: Personal details
    replace_text(paragraph, '<<TITLE>>', form_data['your_title'] or 'N/A')
    replace_text(paragraph, '<<NAME>>', form_data['your_name'] or 'N/A')
    replace_text(paragraph, '<<EMAIL>>', form_data['your_email'] or 'N/A')
    replace_text(paragraph, '<<PHONE>>', form_data['your_phone_number'] or 'N/A')
    replace_text(paragraph, '<<ADDRESS>>', form_data['your_address1'] or 'N/A')
    replace_text(paragraph, '<<POSTCODE>>', form_data['your_postcode'] or 'N/A')

    # Section 10: Communication preferences
    replace_text(paragraph, '<<COMMUNICATION>>', form_data['communication_needs'] or 'N/A')
    replace_text(paragraph, '<<PREFERRED_CONTACT_METHOD>>', form_data['preferred_contact_method'] or 'N/A')

    # Section 11: Declaration
    replace_text(paragraph, '<<DECLARATION>>', "Yes" if form_data['declaration'] else "No")

def replace_judge_placeholders(paragraph, form_data, replace_text):
    replace_text(paragraph, '<<TITLE>>', form_data.get('your_title', 'N/A'))
    replace_text(paragraph, '<<NAME>>', form_data.get('your_name', 'N/A'))
    replace_text(paragraph, '<<ADDRESS>>', "{} {} {}".format(
        form_data.get('your_address1', ''),
        form_data.get('your_address2', ''),
        form_data.get('your_address3', '')
    ).strip())
    replace_text(paragraph, '<<PHONE_NUMBER>>', form_data.get('your_phone_number', 'N/A'))
    replace_text(paragraph, '<<EMAIL>>', form_data.get('your_email', 'N/A'))
    replace_text(paragraph, '<<HEARING_DATE>>', form_data.get('hearing_date', 'N/A'))
    replace_text(paragraph, '<<HEARING_VENUE>>', form_data.get('hearing_venue', 'N/A'))
    replace_text(paragraph, '<<JUDICIAL_OFFICER>>', form_data.get('judicial_officer', 'N/A'))
    replace_text(paragraph, '<<HEARING_NUMBER>>', form_data.get('hearing_number', 'N/A'))
    replace_text(paragraph, '<<MISCONDUCT_DATE>>', form_data.get('misconduct_date', 'N/A'))
    replace_text(paragraph, '<<COMPLAINT>>', form_data.get('complaint_details', 'N/A'))
    replace_text(paragraph, '<<ADDITIONAL_DATES>>', form_data.get('additional_dates', 'N/A'))
    replace_text(paragraph, '<<GUIDANCE_READ>>', 'Yes' if form_data.get('guidance_read') else 'No')
    replace_text(paragraph, '<<UNDERSTANDING_REMIT>>', 'Yes' if form_data.get('understanding_remit') else 'No')

def replace_newspaper_placeholders(paragraph, form_data, replace_text):
    replace_text(paragraph, '<<TITLE>>', form_data.get('your_title', 'N/A'))
    replace_text(paragraph, '<<NAME>>', form_data.get('full_name', 'N/A'))
    replace_text(paragraph, '<<ADDRESS>>', form_data.get('full_address', ''),)
    replace_text(paragraph, '<<PHONE_NUMBER>>', form_data.get('mobile_phone', 'N/A'))
    replace_text(paragraph, '<<EMAIL>>', form_data.get('email_address', 'N/A'))
    replace_text(paragraph, '<<COMPLAIN_ON_BEHALF>>', form_data.get('complain_on_behalf', 'N/A'))
    replace_text(paragraph, '<<RELATIONSHIP_TO_PERSON>>', form_data.get('relationship_to_person', 'N/A'))
    replace_text(paragraph, '<<CONSENT_TO_COMPLAIN>>', form_data.get('consent_to_complain', 'N/A'))
    replace_text(paragraph, '<<REPORTED_PERSON>>', form_data.get('person_or_firm_complained_about', 'N/A'))
    replace_text(paragraph, '<<FIRM_ADDRESS>>', form_data.get('firm_postal_address', 'N/A'))
    replace_text(paragraph, '<<TYPE_OF_WORK>>', form_data.get('type_of_work', 'N/A'))
    replace_text(paragraph, '<<COMPLAINT>>', form_data.get('complaint_details', 'N/A'))
    replace_text(paragraph, '<<AWARENESS_DATE>>', form_data.get('problem_awareness_date', 'N/A'))
    replace_text(paragraph, '<<IMPACT>>', form_data.get('impact_on_you', 'N/A'))
    replace_text(paragraph, '<<COMPLAINED_FIRM>>', form_data.get('complained_to_firm', 'N/A'))
    replace_text(paragraph, '<<COMPLAINT_DATE>>', form_data.get('date_of_complaint', 'N/A'))
    replace_text(paragraph, '<<FIRM_RESPONSE>>', form_data.get('firm_response', 'N/A')) 
    replace_text(paragraph, '<<FIRM_ACTIONS>>', form_data.get('firm_actions', 'N/A'))
    replace_text(paragraph, '<<DESIRED_RESOLUTION>>', form_data.get('desired_resolution', 'N/A')) 
   
   
    
def replace_bank_placeholders(paragraph, form_data, replace_text):
    replace_text(paragraph, '<<TITLE>>', form_data.get('your_title', 'N/A'))
    replace_text(paragraph, '<<NAME>>', form_data.get('full_name', 'N/A'))
    replace_text(paragraph, '<<ADDRESS>>', form_data.get('full_address', ''),)
    replace_text(paragraph, '<<PHONE_NUMBER>>', form_data.get('mobile_phone', 'N/A'))
    replace_text(paragraph, '<<EMAIL>>', form_data.get('email_address', 'N/A'))
    replace_text(paragraph, '<<COMPLAIN_ON_BEHALF>>', form_data.get('complain_on_behalf', 'N/A'))
    replace_text(paragraph, '<<RELATIONSHIP_TO_PERSON>>', form_data.get('relationship_to_person', 'N/A'))
    replace_text(paragraph, '<<CONSENT_TO_COMPLAIN>>', form_data.get('consent_to_complain', 'N/A'))
    replace_text(paragraph, '<<REPORTED_PERSON>>', form_data.get('person_or_firm_complained_about', 'N/A'))
    replace_text(paragraph, '<<FIRM_ADDRESS>>', form_data.get('firm_postal_address', 'N/A'))
    replace_text(paragraph, '<<TYPE_OF_WORK>>', form_data.get('type_of_work', 'N/A'))
    replace_text(paragraph, '<<COMPLAINT>>', form_data.get('complaint_details', 'N/A'))
    replace_text(paragraph, '<<AWARENESS_DATE>>', form_data.get('problem_awareness_date', 'N/A'))
    replace_text(paragraph, '<<IMPACT>>', form_data.get('impact_on_you', 'N/A'))
    replace_text(paragraph, '<<COMPLAINED_FIRM>>', form_data.get('complained_to_firm', 'N/A'))
    replace_text(paragraph, '<<COMPLAINT_DATE>>', form_data.get('date_of_complaint', 'N/A'))
    replace_text(paragraph, '<<FIRM_RESPONSE>>', form_data.get('firm_response', 'N/A')) 
    replace_text(paragraph, '<<FIRM_ACTIONS>>', form_data.get('firm_actions', 'N/A'))
    replace_text(paragraph, '<<DESIRED_RESOLUTION>>', form_data.get('desired_resolution', 'N/A')) 

def solicitor_form(request):
    if request.method == 'POST':
        form = SolicitorForm(request.POST)
        if form.is_valid():
            form_data = form.cleaned_data
            complaint_text = form.cleaned_data['complaint']

            # Instead of classifying, we get the complaint category directly from the user's selection
            complaint_category = 'solicitor'  # Since this is specifically the solicitor form
            
            similar_cases = find_similar_cases(complaint_category, complaint_text)
            breaches = check_for_breaches(complaint_text)

            # Populate the appropriate complaint form
            populated_form_path = populate_complaint_form(complaint_category, form_data, similar_cases)

            # Render the results page, include path to the populated form
            return render(request, 'complaints/complaint_results.html', {
                "complaint_category": complaint_category,
                "similar_cases": similar_cases,
                "breaches": breaches,
                "complaint_text": complaint_text,
                "populated_form_path": populated_form_path,  # Pass the path for the download
            })
    else:
        form = SolicitorForm()

    return render(request, 'complaints/submit_complaint.html', {'form': form})


def download_form(request, file_path):
    """Serve the populated form as a downloadable file."""
    # Get the absolute path to the file in the media directory
    file_path = os.path.join(settings.MEDIA_ROOT, file_path)

    # Serve the document as a downloadable file
    return FileResponse(open(file_path, 'rb'), as_attachment=True, filename=os.path.basename(file_path))


def jcio_form(request):
    if request.method == 'POST':
        form = JCIOForm(request.POST, request.FILES)
        if form.is_valid():
            form_data = form.cleaned_data
            complaint_text = form.cleaned_data['complaint_details']

            # Directly use judge category
            complaint_category = 'judge'

            similar_cases = find_similar_cases(complaint_category, complaint_text)
            breaches = check_for_breaches(complaint_text)
            
            # Populate the appropriate complaint form
            populated_form_path = populate_complaint_form(complaint_category, form_data, similar_cases)
            # Redirect to a success page or provide a download link
            return render(request, 'complaints/complaint_results.html', {
                "complaint_category": complaint_category,
                "similar_cases": similar_cases,
                "breaches": breaches,
                "complaint_text": complaint_text,
                "populated_form_path": populated_form_path,  # Pass the path for the download
            })
    else:
        form = JCIOForm()

    return render(request, 'complaints/jcio_form.html', {'form': form})


def barrister_form(request):
    if request.method == 'POST':
        form = BarristerForm(request.POST, request.FILES)
        if form.is_valid():
            # Create PDF from form data
            form_data = form.cleaned_data
            complaint_text = form.cleaned_data['event_information']
            
            complaint_category = 'barrister'  # Since this is specifically the barrister form

            if complaint_category == "unknown":
                return render(request, 'complaints/error.html', {
                    "error_message": "Complaint category could not be identified."
                })

            similar_cases = find_similar_cases(complaint_category, complaint_text)
            breaches = check_for_breaches(complaint_text)
            # Populate the appropriate complaint form
            populated_form_path = populate_complaint_form(complaint_category, form_data, similar_cases)
            

            # Redirect to a success page or provide a download link
            return render(request, 'complaints/complaint_results.html', {
                "complaint_category": complaint_category,
                "similar_cases": similar_cases,
                "breaches": breaches,
                "complaint_text": complaint_text,
                "populated_form_path": populated_form_path,  # Pass the path for the download
            })  # Change to the correct URL
    else:
        form = BarristerForm()

    return render(request, 'complaints/barrister_form.html', {'form': form})


def success_page(request):
    return render(request, 'complaints/success.html')  #


def newspaper_form(request):
    if request.method == 'POST':
        form = NewspaperForm(request.POST, request.FILES)
        if form.is_valid():
            form_data = form.cleaned_data
            complaint_text = form.cleaned_data['complaint_details']

            # Directly use newspaper category
            complaint_category = 'newspaper'

            similar_cases = find_similar_cases(complaint_category, complaint_text)
            breaches = check_for_breaches(complaint_text)
            
            # Populate the appropriate complaint form
            populated_form_path = populate_complaint_form(complaint_category, form_data, similar_cases)
            # Redirect to a success page or provide a download link
            return render(request, 'complaints/complaint_results.html', {
                "complaint_category": complaint_category,
                "similar_cases": similar_cases,
                "breaches": breaches,
                "complaint_text": complaint_text,
                "populated_form_path": populated_form_path,  # Pass the path for the download
            })
    else:
        form = NewspaperForm()

    return render(request, 'complaints/newspaper_form.html', {'form': form})

def bank_form(request):
    if request.method == 'POST':
        form = BankForm(request.POST, request.FILES)
        if form.is_valid():
            form_data = form.cleaned_data
            complaint_text = form.cleaned_data['complaint_details']

            # Directly use bank category
            complaint_category = 'bank'

            similar_cases = find_similar_cases(complaint_category, complaint_text)
            breaches = check_for_breaches(complaint_text)
            
            # Populate the appropriate complaint form
            populated_form_path = populate_complaint_form(complaint_category, form_data, similar_cases)
            # Redirect to a success page or provide a download link
            return render(request, 'complaints/complaint_results.html', {
                "complaint_category": complaint_category,
                "similar_cases": similar_cases,
                "breaches": breaches,
                "complaint_text": complaint_text,
                "populated_form_path": populated_form_path,  # Pass the path for the download
            })
    else:
        form = BankForm()

    return render(request, 'complaints/bank_form.html', {'form': form})

    
class ContactSubmissionView(APIView):
    def post(self, request):
        serializer = SubmissionSerializer(data=request.data)
        if serializer.is_valid():
            submission = serializer.save()
            
            # Email mapping
            EMAIL_MAPPING = {
                'justiceforsport': 'justiceforsport@gmail.com',
                'litsport': 'engagelitsport@gmail.com',
                'stopitcrew': 'stopitcrew@gmail.com'
            }
            
            # Determine recipient email based on site
            site = submission.site.lower() if submission.site else ''
            recipient_email = next(
                (email for site_key, email in EMAIL_MAPPING.items() if site_key in site),
                'justiceforsport@gmail.com'  # default recipient
            )
            
            # Send email
            send_mail(
                subject=f'New message from {submission.name} on {submission.site}',
                message=f'Name: {submission.name}\nEmail: {submission.email}\nAddress: {submission.address}\nMessage: {submission.message}',
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[recipient_email],
            )

            return Response({'message': 'Success'})
        return Response(serializer.errors, status=400)
    
    


class DocumentUploadView(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request):
        serializer = DocumentSerializer(data=request.data)
        if serializer.is_valid():
            document = serializer.save()
            
            # Process document
            processor = DocumentProcessor()
            try:
                processor.process_document(document)
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            except Exception as e:
                document.delete()  # Clean up if processing fails
                return Response(
                    {'error': f'Document processing failed: {str(e)}'},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class QueryView(APIView):
    def post(self, request):
        serializer = QuerySerializer(data=request.data)
        if serializer.is_valid():
            query = serializer.validated_data['query']
            
            # Get relevant chunks
            processor = QueryProcessor()
            relevant_chunks = processor.get_relevant_chunks(query)
            
            if not relevant_chunks:
                return Response(
                    {'error': 'No relevant information found'},
                    status=status.HTTP_404_NOT_FOUND
                )

            # Query GPT
            context = "\n\n".join(relevant_chunks)
            prompt = f"Use the following document context to answer the query:\n\n{context}\n\nQuery: {query}"
            
            try:
                api_key = settings.OPENAI_API_KEY
                openai.api_key = api_key
                response = openai.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": prompt}]
                )
                answer = response.choices[0].message.content
                return Response({'answer': answer})
            except Exception as e:
                return Response(
                    {'error': f'OpenAI API error: {str(e)}'},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
